import fitz
from docx import Document


def extract_text_from_pdf(file):
    if isinstance(file, str):
        document = fitz.open(file)
    else:
        file.seek(0)
        document = fitz.open(stream=file.read(), filetype="pdf")

    try:
        return "\n".join(page.get_text() for page in document)
    finally:
        document.close()


def extract_text_from_docx(file):
    if not isinstance(file, str):
        file.seek(0)

    document = Document(file)
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n".join(blocks)


def extract_text_from_resume(file, filename):
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file)
    if lower_name.endswith(".docx"):
        return extract_text_from_docx(file)
    return ""
